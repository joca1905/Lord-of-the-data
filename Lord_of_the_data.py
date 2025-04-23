# logo da segunda aba, formatar abnt.

import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document  # Para criar arquivos .docx
from PIL import Image, ImageTk
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelForQuestionAnswering
import bibtexparser
import re
import rispy
import customtkinter
import argostranslate.package
import argostranslate.translate
import zipfile
import customtkinter as ctk
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pybtex.database import parse_file
import seaborn as sns
from collections import Counter
import matplotlib.pyplot as plt
import io
from docx.shared import Inches
# Diretório para salvar o modelo e arquivos processados
MODEL_DIR = "./meu_modelo_summarization"
SALVO_DIR = "./arquivos_salvos"
MODEL_NAME = "Falconsai/text_summarization"

PREDEFINED_QUESTIONS = [
        "What is the main contribution or purpose of this study?",
        "How was the study conducted?",
        "Why was this study conducted?",
        "What are the key findings or results?",
        "What are the main discussions or implications of the findings?"
    ]

PERGUNTAS_PREDEFINIDAS = [
"Qual é a principal contribuição ou propósito deste estudo?",
"Como o estudo foi conduzido?",
"Por que este estudo foi conduzido?",
"Quais são as principais descobertas ou resultados?",
"Quais são as principais discussões ou implicações das descobertas?"
]


# Garantir que os diretórios para salvar arquivos existem
os.makedirs(SALVO_DIR, exist_ok=True)


# Função para carregar ou baixar o modelo de sumarização
def load_model():
    if not os.path.exists(MODEL_DIR):
        print("Baixando o modelo... Isso pode levar alguns minutos.")
        tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
        model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)
        tokenizer.save_pretrained(MODEL_DIR)
        model.save_pretrained(MODEL_DIR)
        print(f"Modelo salvo em {MODEL_DIR}")
    else:
        print(f"Carregando o modelo salvo em {MODEL_DIR}...")

    tokenizer = AutoTokenizer.from_pretrained(MODEL_DIR)
    model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_DIR)
    return pipeline("summarization", model=model, tokenizer=tokenizer)


# Inicializa o pipeline de sumarização
summarizer = load_model()


# Diretório para salvar o modelo 1 qa
QA_MODEL_DIR = "./meu_modelo_question_answering"
QA_MODEL_NAME = "distilbert/distilbert-base-uncased-distilled-squad"

# Função para carregar ou baixar o modelo de perguntas e respostas
def load_qa_model():
    if not os.path.exists(QA_MODEL_DIR):
        print("Baixando o modelo de perguntas e respostas... Isso pode levar alguns minutos.")
        tokenizer = AutoTokenizer.from_pretrained(QA_MODEL_NAME)
        model = AutoModelForQuestionAnswering.from_pretrained(QA_MODEL_NAME)
        tokenizer.save_pretrained(QA_MODEL_DIR)
        model.save_pretrained(QA_MODEL_DIR)
        print(f"Modelo salvo em {QA_MODEL_DIR}")
    else:
        print(f"Carregando o modelo salvo em {QA_MODEL_DIR}...")

    tokenizer = AutoTokenizer.from_pretrained(QA_MODEL_DIR)
    model = AutoModelForQuestionAnswering.from_pretrained(QA_MODEL_DIR)
    return pipeline("question-answering", model=model, tokenizer=tokenizer)

# Inicializa o pipeline de perguntas e respostas
question_answerer = load_qa_model()



# Função para carregar e processar arquivos BibTeX ou RIS
def load_file(input_text):
    file_path = filedialog.askopenfilename(filetypes=[("BibTeX files", "*.bib"), ("RIS files", "*.ris")])
    if not file_path:
        return

    try:
        entries = []
        if file_path.endswith(".bib"):
            with open(file_path, "r", encoding="utf-8") as bibtex_file:
                bib_database = bibtexparser.load(bibtex_file)
                entries = bib_database.entries
        elif file_path.endswith(".ris"):
            with open(file_path, "r", encoding="utf-8") as ris_file:
                entries = rispy.load(ris_file)
        else:
            raise ValueError("Formato de arquivo não suportado.")

        # Exibe as informações principais no campo de entrada
        input_text.delete("1.0", tk.END)
        for entry in entries:
            title = entry.get("title", "Título não encontrado")
            authors = entry.get("author", "Autores não encontrados")
            year = entry.get("year", "Ano não encontrado")
            doi = entry.get("doi", "DOI não encontrado")

            article_info = f"Título: {title}\nAutores: {authors}\nAno: {year}\nDOI: {doi}\n\n"
            input_text.insert(tk.END, article_info)

        # Salva o conteúdo processado em um arquivo local
        saved_file_path = os.path.join(SALVO_DIR, os.path.basename(file_path))
        with open(saved_file_path, "w", encoding="utf-8") as output_file:
            output_file.write(input_text.get("1.0", tk.END))
        input_text.insert(tk.END, f"\nInformações salvas em: {saved_file_path}")

    except Exception as e:
        input_text.delete("1.0", tk.END)
        input_text.insert(tk.END, f"Erro ao carregar o arquivo: {str(e)}")

# Função para gerar o resumo
def generate_summary(summary_input_text, summary_output_text):
    article = summary_input_text.get("1.0", tk.END).strip()
    if article:
        try:
            summary = summarizer(article, max_length=1000, min_length=30, do_sample=False)[0]["summary_text"]
            summary_output_text.delete("1.0", tk.END)
            summary_output_text.insert(tk.END, summary)
        except Exception as e:
            summary_output_text.delete("1.0", tk.END)
            summary_output_text.insert(tk.END, f"Erro ao gerar o resumo: {str(e)}")
    else:
        summary_output_text.delete("1.0", tk.END)
        summary_output_text.insert(tk.END, "Por favor, insira um texto para resumir.")


# Defina translate_text como uma função independente
def translate_text_model(text_to_translate, from_lang, to_lang):
    """Traduz o texto de um idioma para outro usando Argos Translate."""
    try:
        # Atualiza os pacotes disponíveis
        argostranslate.package.update_package_index()
        available_packages = argostranslate.package.get_available_packages()

        # Tradução direta
        package_to_install = next(
            (pkg for pkg in available_packages if pkg.from_code == from_lang and pkg.to_code == to_lang),
            None
        )

        if package_to_install:
            argostranslate.package.install_from_path(package_to_install.download())
            translated_text = argostranslate.translate.translate(text_to_translate, from_lang, to_lang)
            return translated_text
        else:
            messagebox.showwarning("Aviso", f"Nenhum pacote disponível para traduzir de {from_lang} para {to_lang}.")
            return None

    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao traduzir o texto: {str(e)}")
        return None



# Atualizando a classe App para incluir a aba de tradução
class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        #self.geometry("1000x700")
        self.title("Lord of the Data")

        icon_path = "C:/Users/jocas/Downloads/data_of_data.ico"  # Substitua pelo caminho correto
        self.iconbitmap(icon_path)


        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        self.tabview = customtkinter.CTkTabview(self, width=250)
        self.tabview.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")

        # Criando as abas existentes
        self.tabview.add("Home")
        self.tabview.add("Configurações")
        self.tabview.add("Revisão Sistemática")
        self.tabview.add("Execução")
        self.tabview.add("Sumarização")
        self.tabview.add("Perguntas e Respostas")
        self.tabview.add("Tradução")  # Nova aba


        # Chamando as funções para preencher as abas
        self.create_home_tab(self.tabview.tab("Home"))
        self.create_settings_tab(self.tabview.tab("Configurações"))
        self.create_review_tab(self.tabview.tab("Revisão Sistemática"))
        self.create_execution_tab(self.tabview.tab("Execução"))
        self.create_summary_tab(self.tabview.tab("Sumarização"))
        self.create_translation_tab(self.tabview.tab("Tradução"))  # Nova aba
        self.create_qa_tab(self.tabview.tab("Perguntas e Respostas"))

        self.metadata = []
        self.metadata_dados = []
        self.metadata_ref = []
        self.formatted_qa = []
        self.decisions_data = []

    def create_home_tab(self, tab):
        """Cria os elementos para a aba inicial (home)."""
        # Configuração do grid para centralizar os elementos
        tab.grid_rowconfigure((0, 1, 2, 3, 4, 5), weight=1)
        tab.grid_columnconfigure(0, weight=1)

        # Logo do programa
        logo_image = customtkinter.CTkImage(light_image=Image.open("C:/Users/jocas/Downloads/data_of_data.ico"),
                                            dark_image=Image.open("C:/Users/jocas/Downloads/data_of_data.ico"),
                                            size=(300, 300))  # Ajuste o tamanho conforme necessário
        logo_label = customtkinter.CTkLabel(tab, image=logo_image, text="")
        logo_label.grid(row=0, column=0, pady=(20, 10), sticky="n")

        # Nome do programa
        program_name = customtkinter.CTkLabel(tab, 
                                            text="Lord of the data", 
                                            font=customtkinter.CTkFont(size=24, weight="bold"))
        program_name.grid(row=1, column=0, pady=(0, 10), sticky="n")

        # Slogan ou descrição
        slogan = customtkinter.CTkLabel(tab, 
                                        text="Sua ferramenta para revisão sistemática de artigos científicos.", 
                                        font=customtkinter.CTkFont(size=14))
        slogan.grid(row=2, column=0, pady=(0, 20), sticky="n")

        # Rodapé com informações adicionais
        footer = customtkinter.CTkLabel(tab, 
                                        text="Versão 0.0001 | © 2025 Joaquim Osterwald Frota Moura Filho", 
                                        font=customtkinter.CTkFont(size=12))
        footer.grid(row=5, column=0, pady=(20, 10), sticky="s")

    def create_settings_tab(self, tab):
        """Cria os elementos para a aba de configurações."""
        tab.grid_rowconfigure((0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16), weight=1)
        tab.grid_columnconfigure((0, 1), weight=1)

        # Botões de salvar e carregar
        save_button = customtkinter.CTkButton(tab, text="Salvar Revisão", command=self.save_review)
        save_button.grid(row=0, column=0, padx=10, pady=10, sticky="ew")

        load_button = customtkinter.CTkButton(tab, text="Carregar Revisão", command=self.load_review)
        load_button.grid(row=0, column=1, padx=10, pady=10, sticky="ew")

        # Menu de seleção de formato
        format_label = customtkinter.CTkLabel(tab, text="Formato do arquivo:")
        format_label.grid(row=2, column=0, padx=20, pady=(10, 5), sticky="e")

        # Adicione uma variável para armazenar a escolha do usuário
        self.format_option = customtkinter.StringVar(value="Docx")  # Valor padrão

        format_option = customtkinter.CTkOptionMenu(tab, values=["Docx", "Latex"], variable=self.format_option)
        format_option.grid(row=2, column=1, padx=10, pady=10, sticky="w")
        format_option.set("Docx")  # Valor padrão

        # Seleção do modelo de question-answering
        #model_label = customtkinter.CTkLabel(tab, text="Modelo de QA:")
        #model_label.grid(row=5, column=0, padx=20, pady=(10, 5), sticky="e")
        
        # Adicione uma variável para armazenar a escolha do usuário
        #self.model_var = customtkinter.StringVar(value="distilbert-base-uncased-distilled-squad")  # Valor padrão

        # Lista de modelos disponíveis
        #model_options = [
        #    "distilbert-base-uncased-distilled-squad",  # Modelo 1
        #    "DeepSeek-R1-Distill-Qwen-1.5B",    # Modelo 2
        #]

        #model_option = customtkinter.CTkOptionMenu(tab, 
        #                                        values=model_options, 
        #                                        variable=self.model_var)  # Associa a variável
        #model_option.grid(row=5, column=1, padx=10, pady=(10, 5), sticky="w")

        # Seleção do modo de aparência
        appearance_label = customtkinter.CTkLabel(tab, text="Modo de Aparência:")
        appearance_label.grid(row=3, column=0, padx=20, pady=(10, 5), sticky="e")
        
        appearance_option = customtkinter.CTkOptionMenu(tab, 
                                                        values=["Light", "Dark"], 
                                                        command=customtkinter.set_appearance_mode)
        appearance_option.grid(row=3, column=1, padx=10, pady=(10, 5), sticky="w")
        appearance_option.set("Light")  # Configuração padrão

        
        reference_type = customtkinter.CTkLabel(tab, text="Tipo de referência:")
        reference_type.grid(row=4, column=0, padx=20, pady=(10, 5), sticky="e")
        
        # Adicione uma variável para armazenar a escolha do usuário
        self.reference_var = customtkinter.StringVar(value="ABNT")  # Valor padrão

        reference_option = customtkinter.CTkOptionMenu(tab, 
                                                        values=["ABNT", "Havard"], 
                                                        command=self.update_style,  # Chama função ao mudar
                                                        variable=self.reference_var)  # Associa a variável
        reference_option.grid(row=4, column=1, padx=10, pady=(10, 5), sticky="w")



    def save_review(self):
        """Salva a revisão no formato selecionado."""
        selected_format = self.format_option.get()  # Obtém o formato selecionado

        if selected_format == "Docx":
            self.save_to_docx()
        elif selected_format == "Latex":
            self.save_to_latex()
        else:
            messagebox.showerror("Erro", "Formato de salvamento inválido.")

    def load_review(self):
        """Carrega a revisão no formato selecionado."""
        selected_format = self.format_option.get()  # Obtém o formato selecionado

        if selected_format == "Docx":
            self.load_from_docx()
        elif selected_format == "Latex":
            self.load_from_latex()
        else:
            messagebox.showerror("Erro", "Formato de carregamento inválido.")
    
    def update_style(self, choice):
        """Atualiza a variável de estilo quando o usuário faz uma escolha."""
        self.selected_style = self.reference_var.get()  # Atualiza o estilo com o valor da StringVar

    def create_review_tab(self, tab):
        """Cria os elementos para a aba de informações da revisão."""
        tab.grid_rowconfigure((0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15), weight=1)  # Expandindo configuração
        tab.grid_columnconfigure((1, 2), weight=1)

        # Título da revisão
        title_label = customtkinter.CTkLabel(tab, text="Título da Revisão:")
        title_label.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="e")
        self.title_entry = customtkinter.CTkEntry(tab, placeholder_text="Digite o título aqui")
        self.title_entry.grid(row=0, column=1, padx=10, pady=(10, 5), sticky="ew")

        # Pesquisadores
        researchers_label = customtkinter.CTkLabel(tab, text="Pesquisadores:")
        researchers_label.grid(row=1, column=0, padx=10, pady=(5, 5), sticky="e")
        self.researchers_entry = customtkinter.CTkEntry(tab, placeholder_text="Digite os nomes dos pesquisadores")
        self.researchers_entry.grid(row=1, column=1, padx=10, pady=(5, 5), sticky="ew")

        # Descrição
        description_label = customtkinter.CTkLabel(tab, text="Descrição:")
        description_label.grid(row=2, column=0, padx=10, pady=(5, 10), sticky="ne")
        self.description_textbox = customtkinter.CTkTextbox(tab, height=50)
        self.description_textbox.grid(row=2, column=1, padx=10, pady=(5, 10), sticky="nsew")

        # Objetivo
        objective_label = customtkinter.CTkLabel(tab, text="Objetivo:")
        objective_label.grid(row=3, column=0, padx=10, pady=(10, 5), sticky="e")
        self.objective_entry = customtkinter.CTkEntry(tab, placeholder_text="Defina o objetivo da revisão")
        self.objective_entry.grid(row=3, column=1, padx=10, pady=(10, 5), sticky="ew")

        # Linguagem dos artigos
        languages_label = customtkinter.CTkLabel(tab, text="Linguagem dos artigos:")
        languages_label.grid(row=4, column=0, padx=10, pady=(5, 5), sticky="ne")
        
        self.languages_selected = []
        self.language_options = ["Português", "Inglês", "Espanhol", "Russo", "Chinês", "Outro"]
        self.language_checkboxes = []
        columns = 2
        for i, language in enumerate(self.language_options):
            var = customtkinter.StringVar(value="")
            checkbox = customtkinter.CTkCheckBox(tab, text=language, variable=var, onvalue=language, offvalue="")
            row, col = 5 + i // columns, i % columns
            checkbox.grid(row=row, column=col, padx=10, pady=2, sticky="w")
            self.language_checkboxes.append((checkbox, var))
        

        # Editoras
        publishers_label = customtkinter.CTkLabel(tab, text="Editoras:")
        publishers_label.grid(row=8, column=0, padx=10, pady=(5, 5), sticky="ne")
        
        self.publishers_selected = []
        self.publisher_options = ["Springer", "Elsevier", "IEEE", "MDPI", "Wiley", "Outra"]
        self.publisher_checkboxes = []
        for i, publisher in enumerate(self.publisher_options):
            var = customtkinter.StringVar(value="")
            checkbox = customtkinter.CTkCheckBox(tab, text=publisher, variable=var, onvalue=publisher, offvalue="")
            row, col = 9 + i // columns, i % columns
            checkbox.grid(row=row, column=col, padx=10, pady=2, sticky="w")
            self.publisher_checkboxes.append((checkbox, var))

        # Palavras-chave
        keywords_label = customtkinter.CTkLabel(tab, text="Palavras-chave:")
        keywords_label.grid(row=12, column=0, padx=10, pady=(5, 10), sticky="ne")
        self.keywords_textbox = customtkinter.CTkEntry(tab, placeholder_text="Defina as palavras-chave da revisão")
        self.keywords_textbox.grid(row=12, column=1, padx=10, pady=(5, 20), sticky="nsew")

        # strings de pesquisa
        string_label = customtkinter.CTkLabel(tab, text="Strings de pesquisa:")
        string_label.grid(row=13, column=0, padx=10, pady=(5, 10), sticky="ne")
        self.strings_textbox = customtkinter.CTkEntry(tab, placeholder_text="Defina as strings de pesquisa da revisão")
        self.strings_textbox.grid(row=13, column=1, padx=10, pady=(5, 20), sticky="nsew")

        # Critérios de seleção
        criteria_selection = customtkinter.CTkLabel(tab, text="Critérios de seleção:")
        criteria_selection.grid(row=14, column=0, padx=10, pady=(5, 10), sticky="ne")
        self.criteria_textbox_s = customtkinter.CTkEntry(tab, placeholder_text="Defina os critérios de seleção")
        self.criteria_textbox_s.grid(row=14, column=1, padx=10, pady=(5, 20), sticky="nsew")

        # Critérios de exclusão
        criteria_exclusion = customtkinter.CTkLabel(tab, text="Critérios de exclusão:")
        criteria_exclusion.grid(row=15, column=0, padx=10, pady=(5, 10), sticky="ne")
        self.criteria_textbox_e = customtkinter.CTkEntry(tab, placeholder_text="Defina os critérios de exclusão")
        self.criteria_textbox_e.grid(row=15, column=1, padx=10, pady=(5, 20), sticky="nsew")
    
    def format_references(self, metadata):
        """Formata as referências de acordo com o estilo selecionado (ABNT ou Harvard)."""
        formatted_references = []
        
        for entry in metadata:
            title = entry.get('Title', 'N/A')
            authors = entry.get('Authors', ['N/A'])
            journal = entry.get('Journal', 'N/A')
            volume = entry.get('Volume', '')
            number = entry.get('Number', '')
            pages = entry.get('Pages', '')
            doi = entry.get('DOI', 'N/A')
            year = entry.get('Year', 'N/A')  # Obtém o ano
            
            # Formatação dos autores no estilo ABNT
            if len(authors) > 3:
                formatted_authors = f"{authors[0].split()[-1].upper()}, {authors[0].split()[0]} et al."
            else:
                formatted_authors = ', '.join(
                    [f"{author.split()[-1].upper()}, {' '.join(author.split()[0:-1])}" for author in authors]
                )
            
            selected_style = self.reference_var.get()  # Acessa diretamente a StringVar
            
            if selected_style == "ABNT":
                # Formatação no estilo ABNT
                journal_info = f"**{journal}**"
                if volume:
                    journal_info += f", {volume}"
                if number:
                    journal_info += f"({number})"
                if pages:
                    journal_info += f", {pages}"
                
                formatted_entry = (
                    f"{formatted_authors}. {title}. {journal_info}, {year}. DOI: {doi}."
                )
            elif selected_style == "Harvard":
                # Formatação no estilo Harvard
                formatted_entry = (
                    f"{', '.join(authors)} ({year}). {title}. **{journal}**, {volume}({number}), {pages}. DOI: {doi}."
                )
            else:
                # Formato padrão (caso o estilo não seja ABNT ou Harvard)
                formatted_entry = (
                    f"{', '.join(authors)}. {title}. **{journal}**, {volume}({number}), {pages}, {year}. DOI: {doi}."
                )
            
            # Adiciona a referência formatada à lista
            formatted_references.append(formatted_entry)
        
        return formatted_references  # Retorna uma lista de referências formatadas
   
    
    def create_execution_tab(self, tab):
        """Cria os elementos para a aba de execução."""
        tab.grid_rowconfigure(0, weight=0)  # Linha do botão não se expande
        tab.grid_rowconfigure(1, weight=1)  # Linha do campo de texto se expande
        tab.grid_columnconfigure(0, weight=1)  # Coluna se expande

        # Botão para carregar arquivos
        load_button = customtkinter.CTkButton(tab, text="Carregar Arquivos RIS/BibTeX",
                                               command=self.load_files)
        load_button.grid(row=0, column=0, padx=10, pady=10, sticky="ew")  # Mudar para "ew" para expandir na horizontal

        # Campo de texto para exibir informações
        self.input_text = customtkinter.CTkTextbox(tab, wrap="word")
        self.input_text.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")  # "nsew" para expandir em todas as direções


    def load_files(self):
        """Carrega arquivos RIS ou BibTeX, extrai metadados e processa os abstracts."""
        file_paths = filedialog.askopenfilenames(
            title="Selecionar Arquivos",
            filetypes=[("Arquivo RIS", "*.ris"), ("Arquivo BibTeX", "*.bib")]
        )
        if file_paths:
            self.input_text.delete('1.0', customtkinter.END)  # Limpa o campo de texto

            all_metadata = []  # Armazena todos os metadados extraídos
            for file_path in file_paths:
                content = self.extract_metadata(file_path)
                self.input_text.insert(customtkinter.END, f"Metadados de {file_path}:\n{content}\n\n")
                all_metadata.append(content)  # Adiciona os metadados extraídos à lista

            self.metadata = "\n\n".join(all_metadata)

            # Extrai os metadados limpos (dicionários)
            self.clean_metadata = []
            for file_path in file_paths:
                self.clean_metadata.extend(self.extract_clean_metadata(file_path))

            # Processa os abstracts e responde as perguntas pré-definidas
            results = self.process_abstracts(self.clean_metadata)

            # Exibe os resultados no campo de texto
            self.input_text.insert(customtkinter.END, "Respostas geradas:\n\n")
            for article, answers in results.items():
                self.input_text.insert(customtkinter.END, f"{article}:\n")
                for question, answer in answers.items():
                    # Extrai o score e a resposta
                    score = answer.get('score', 0)  # Obtém o score (padrão 0 se não existir)
                    response = answer.get('answer', 'N/A')  # Obtém a resposta (padrão 'N/A' se não existir)
                    
                    # Converte o score para porcentagem (multiplicando por 100 e arredondando)
                    score_percent = round(score * 100, 2)  # Arredonda para 2 casas decimais
                    
                    # Exibe o score em porcentagem e a resposta
                    self.input_text.insert(customtkinter.END, f"- {question}\n")
                    self.input_text.insert(customtkinter.END, f"  Score: {score_percent}%\n")
                    self.input_text.insert(customtkinter.END, f"  Resposta: {response}\n\n")

            # Formata as perguntas e respostas
            self.formatted_qa = self.format_qa_results(results)

            self.metadata_dados = self.clean_metadata

            # Formata as referências com os dados extraídos
            self.metadata_ref = self.format_references(self.clean_metadata)

            # Exibe as referências formatadas no campo de texto
            self.input_text.insert(customtkinter.END, "Referências formatadas:\n\n")
            # Junta as referências em uma única string com quebras de linha
            formatted_references_text = "\n\n".join(self.metadata_ref)
            self.input_text.insert(customtkinter.END, formatted_references_text)

            # Coleta as decisões do usuário sobre os artigos
            self.collect_decisions()


    def extract_clean_metadata(self, file_path):
        """Extrai metadados de um arquivo BibTeX ou RIS e retorna como uma lista de dicionários."""
        metadata_list = []

        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            is_bibtex = content.strip().startswith('@')  # Verifica se é um arquivo BibTeX

        if is_bibtex:
            # Processamento correto do BibTeX
            bib_data = parse_file(file_path)
            for key, entry in bib_data.entries.items():
                metadata = {
                    'Title': entry.fields.get('title', '').replace('{', '').replace('}', ''),
                    'Year': entry.fields.get('year', ''),
                    'Journal': entry.fields.get('journal', ''),
                    'DOI': entry.fields.get('doi', ''),
                    'Authors': [' '.join(person.first_names + person.last_names) for person in entry.persons.get('author', [])],
                    'Abstract': entry.fields.get('abstract', ''),
                    'Volume': entry.fields.get('volume', ''),  # Extrai o volume
                    'Number': entry.fields.get('number', ''),  # Extrai o número
                    'Pages': entry.fields.get('pages', '')    # Extrai as páginas
                }
                metadata_list.append(metadata)
        else:
            # Processamento do RIS (mantendo sua lógica original)
            current_metadata = {}
            with open(file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if line.startswith('TI  - ') or line.startswith('T1  - '):  # Título
                        if current_metadata:  # Adiciona metadados anteriores à lista
                            metadata_list.append(current_metadata)
                            current_metadata = {}
                        current_metadata['Title'] = line[6:].strip()
                    elif line.startswith('AU  - '):  # Autor
                        if 'Authors' not in current_metadata:
                            current_metadata['Authors'] = []
                        current_metadata['Authors'].append(line[6:].strip())
                    elif line.startswith('JO  - '):  # Revista
                        current_metadata['Journal'] = line[6:].strip()
                    elif line.startswith('DO  - '):  # DOI
                        current_metadata['DOI'] = line[6:].strip()
                    elif line.startswith('AB  - '):  # Resumo
                        current_metadata['Abstract'] = line[6:].strip()
                    elif line.startswith('PY  - '):  # Ano de publicação
                        current_metadata['Year'] = line[6:].strip()
                    elif line.startswith('VL  - '):  # Volume
                        current_metadata['Volume'] = line[6:].strip()
                    elif line.startswith('IS  - '):  # Número (Issue)
                        current_metadata['Number'] = line[6:].strip()
                    elif line.startswith('SP  - '):  # Página inicial
                        current_metadata['Pages'] = line[6:].strip()
                    elif line.startswith('EP  - '):  # Página final
                        if 'Pages' in current_metadata:
                            current_metadata['Pages'] += f"-{line[6:].strip()}"  # Concatena com a página inicial
                        else:
                            current_metadata['Pages'] = line[6:].strip()
                    elif line.startswith('ER  - '):  # Fim do segmento RIS
                        if current_metadata:  # Adiciona a entrada antes de resetar
                            metadata_list.append(current_metadata)
                        current_metadata = {}

            # Adiciona o último conjunto de metadados, se existir
            if current_metadata:
                metadata_list.append(current_metadata)

        return metadata_list


    def extract_metadata(self, file_path):
        """Extrai metadados de um arquivo RIS ou BibTeX."""
        metadata_list = []  # Usado para armazenar metadados de múltiplos segmentos
        current_metadata = {}
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            is_bibtex = content.startswith('@')  # Verifica se é um arquivo BibTeX
            file.seek(0)  # Reseta o cursor do arquivo para leitura

            if is_bibtex:
                # Processamento correto do BibTeX usando pybtex
                bib_data = parse_file(file_path)
                for key, entry in bib_data.entries.items():
                    metadata = {
                        'Title': entry.fields.get('title', '').replace('{', '').replace('}', ''),
                        'Authors': [' '.join(person.first_names + person.last_names) for person in entry.persons.get('author', [])],
                        'Abstract': entry.fields.get('abstract', ''),
                        'Year': entry.fields.get('year', ''),
                        'Journal': entry.fields.get('journal', ''),
                        'DOI': entry.fields.get('doi', '')
                        
                    }
                        
                    metadata_list.append(metadata)

            else:
                # Lógica para arquivos RIS
                for line in file:
                    if line.startswith('TI  - ') or line.startswith('T1  - '):  # Título
                        if current_metadata:  # Se houver metadados anteriores, adiciona à lista
                            metadata_list.append(current_metadata)
                            current_metadata = {}
                        current_metadata['Title'] = line[6:].strip()  # Extraindo o título
                    elif line.startswith('AU  - '):  # Autor
                        if 'Authors' not in current_metadata:
                            current_metadata['Authors'] = []
                        current_metadata['Authors'].append(line[6:].strip())
                    elif line.startswith('JO  - '):  # Revista
                        current_metadata['Journal'] = line[6:].strip()
                    elif line.startswith('DO  - '):  # DOI
                        current_metadata['DOI'] = line[6:].strip()
                    elif line.startswith('AB  - '):  # Resumo
                        current_metadata['Abstract'] = line[6:].strip()
                    elif line.startswith('PY  - '):  # Ano de publicação em RIS
                        current_metadata['Year'] = line[6:].strip()  # Extraindo o ano
                    elif line.startswith('ER  - '):  # Fim do segmento RIS
                        if current_metadata:  # Para adicionar o último segmento
                            metadata_list.append(current_metadata)
                        current_metadata = {}

                # Adiciona o último conjunto de metadados, se existir
                if current_metadata:
                    metadata_list.append(current_metadata)
            
        # Formatação da saída
        formatted_output = []
        for metadata in metadata_list:
            output = []
            output.append(f"Título: {metadata.get('Title', 'N/A')}")
            if 'Authors' in metadata:
                output.append(f"Autores: {', '.join(metadata['Authors'])}")
            output.append(f"Revista: {metadata.get('Journal', 'N/A')}")
            output.append(f"DOI: {metadata.get('DOI', 'N/A')}")
            if 'Abstract' in metadata:
                output.append(f"Resumo: {metadata.get('Abstract', 'N/A')}")
            if 'Year' in metadata:
                output.append(f"Ano: {metadata.get('Year', 'N/A')}")
            output.append("")  # Para adicionar uma linha em branco entre os conjuntos

            formatted_output.append('\n'.join(output))
        
        return '\n\n'.join(formatted_output).strip()  # Retorna todos os metadados formatados
    
    def collect_decisions(self):
        """Coleta as decisões do usuário sobre os artigos."""
        if not hasattr(self, 'clean_metadata') or not self.clean_metadata:
            print("Nenhum artigo carregado.")
            return

        # Cria uma nova janela para coletar as decisões
        decision_window = ctk.CTkToplevel(self)
        decision_window.title("Seleção dos artigos")
        decision_window.geometry("800x600")

        # Tenta carregar o ícone da janela
        icon_path = "C:/Users/jocas/Downloads/data_of_data.ico"  # Substitua pelo caminho correto
        try:
            self.iconbitmap(icon_path)
        except Exception as e:
            print(f"Erro ao carregar o ícone: {e}")

        # Frame para os artigos
        frame = ctk.CTkScrollableFrame(decision_window)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Lista para armazenar as decisões
        self.articles_decision = []

        # Coleta as decisões para cada artigo
        for i, article in enumerate(self.clean_metadata):
            # Título do artigo
            title_label = ctk.CTkLabel(frame, text=f"Artigo {i+1}: {article.get('Title', 'N/A')}", font=("Helvetica", 14, "bold"))
            title_label.pack(pady=5, anchor="w")

            # Aceitar ou rejeitar artigo
            decision_var = ctk.StringVar(value="Escolha")  # Valor padrão
            decision_label = ctk.CTkLabel(frame, text="Decisão:")
            decision_label.pack(pady=5, anchor="w")
            decision_option = ctk.CTkOptionMenu(
                frame,
                values=["Aceitar", "Rejeitar"],
                variable=decision_var
            )
            decision_option.pack(pady=5, anchor="w")

            # Critério de seleção (só habilitado se o artigo for aceito)
            selection_label = ctk.CTkLabel(frame, text="Critério de seleção:")
            selection_label.pack(pady=5, anchor="w")
            selection_entry = ctk.CTkEntry(frame, width=600)
            selection_entry.pack(pady=5, anchor="w")
            selection_entry.configure(state="disabled")  # Inicialmente desabilitado

            # Critério de exclusão (só habilitado se o artigo for rejeitado)
            exclusion_label = ctk.CTkLabel(frame, text="Critério de exclusão:")
            exclusion_label.pack(pady=5, anchor="w")
            exclusion_entry = ctk.CTkEntry(frame, width=600)
            exclusion_entry.pack(pady=5, anchor="w")
            exclusion_entry.configure(state="disabled")  # Inicialmente desabilitado

            # Prioridade de leitura (só habilitado se o artigo for aceito)
            priority_var = ctk.StringVar(value="Média")  # Valor padrão
            priority_label = ctk.CTkLabel(frame, text="Prioridade de leitura:")
            priority_label.pack(pady=5, anchor="w")
            priority_option = ctk.CTkOptionMenu(
                frame,
                values=["Alta", "Média", "Baixa"],
                variable=priority_var
            )
            priority_option.pack(pady=5, anchor="w")
            priority_option.configure(state="disabled")  # Inicialmente desabilitado

            # Função para habilitar/desabilitar campos com base na decisão
            def toggle_fields(decision_var, selection_entry, exclusion_entry, priority_option, priority_var):
                if decision_var.get() == "Aceitar":
                    selection_entry.configure(state="normal")
                    exclusion_entry.configure(state="disabled")
                    priority_option.configure(state="normal")
                else:
                    selection_entry.configure(state="disabled")
                    exclusion_entry.configure(state="normal")
                    priority_option.configure(state="disabled")
                    priority_var.set("N/A")  # Reseta a prioridade para "Não aplicável"

            # Atualiza o estado dos campos quando a decisão muda
            decision_var.trace_add("write", lambda *args, dv=decision_var, se=selection_entry, ee=exclusion_entry, po=priority_option, pv=priority_var: toggle_fields(dv, se, ee, po, pv))

            # Armazena os widgets para uso posterior
            self.articles_decision.append({
                'selection_entry': selection_entry,
                'exclusion_entry': exclusion_entry,
                'decision_var': decision_var,
                'priority_var': priority_var
            })

        # Função para coletar os valores antes de fechar a janela
        def confirm_decisions():
            # Armazena os valores dos widgets antes de destruir a janela
            self.decisions_data = []
            for decision in self.articles_decision:
                decision_data = {
                    'decision': decision['decision_var'].get(),
                    'selection': decision['selection_entry'].get() if decision['decision_var'].get() == "Aceitar" else None,
                    'exclusion': decision['exclusion_entry'].get() if decision['decision_var'].get() == "Rejeitar" else None,
                    'priority': decision['priority_var'].get() if decision['decision_var'].get() == "Aceitar" else "Não aplicável"
                }
                self.decisions_data.append(decision_data)
            decision_window.destroy()

        # Botão para confirmar as decisões
        confirm_button = ctk.CTkButton(decision_window, text="Confirmar", command=confirm_decisions)
        confirm_button.pack(pady=10)
    
    
    def create_summary_tab(self, tab):
        """Cria os elementos para a aba de sumarização."""
        tab.grid_rowconfigure((0, 1, 2, 3, 4,5,6,7), weight=1)
        tab.grid_columnconfigure(0, weight=1)

        # Campo de entrada para o texto a ser resumido
        input_label = customtkinter.CTkLabel(tab, text="Texto para resumir:")
        input_label.grid(row=0, column=0, padx=10, pady=(10, 5), sticky="w")
        self.summary_input_text = customtkinter.CTkTextbox(tab, height=150)
        self.summary_input_text.grid(row=1, column=0, padx=10, pady=(5, 10), sticky="nsew")

        # Botão para gerar resumo
        generate_button = customtkinter.CTkButton(tab, text="Gerar Resumo", command=lambda: generate_summary(
            self.summary_input_text, self.summary_output_text))
        generate_button.grid(row=2, column=0, padx=10, pady=5, sticky="ew")

        # Campo de saída para exibir o resumo gerado
        output_label = customtkinter.CTkLabel(tab, text="Resumo gerado:")
        output_label.grid(row=3, column=0, padx=10, pady=(10, 5), sticky="w")
        self.summary_output_text = customtkinter.CTkTextbox(tab, height=150)
        self.summary_output_text.grid(row=4, column=0, padx=10, pady=(5, 10), sticky="nsew")

    def create_qa_tab(self, tab):
        """Cria os elementos para a aba de perguntas e respostas."""
        tab.grid_rowconfigure((0, 1, 2, 3, 4, 5), weight=1)
        tab.grid_columnconfigure((0, 1), weight=1)

        # Entrada para o contexto
        context_label = customtkinter.CTkLabel(tab, text="Contexto:")
        context_label.grid(row=0, column=0, padx=20, pady=10, sticky="e")
        context_text = customtkinter.CTkTextbox(tab, height=200)
        context_text.grid(row=0, column=1, padx=20, pady=10, sticky="nsew")

        # Entrada para a pergunta
        question_label = customtkinter.CTkLabel(tab, text="Pergunta:")
        question_label.grid(row=1, column=0, padx=20, pady=10, sticky="e")
        question_text = customtkinter.CTkTextbox(tab, height=50)
        question_text.grid(row=1, column=1, padx=20, pady=10, sticky="nsew")

        # Campo para a resposta
        answer_label = customtkinter.CTkLabel(tab, text="Resposta:")
        answer_label.grid(row=2, column=0, padx=20, pady=10, sticky="e")
        answer_text = customtkinter.CTkTextbox(tab, height=50)
        answer_text.grid(row=2, column=1, padx=20, pady=10, sticky="nsew")

        # Função para processar a pergunta e gerar a resposta
        def answer_question(context_text, question_text, output_text=None):
            if context_text.strip() and question_text.strip():
                try:
                    result = question_answerer(question=question_text, context=context_text)
                    if output_text:  # Se um campo de texto for fornecido, exibe a resposta nele
                        output_text.delete("1.0", tk.END)
                        output_text.insert(tk.END, result["answer"])
                    return result["answer"]
                except Exception as e:
                    return f"Erro ao processar a pergunta: {str(e)}"
            else:
                return "Contexto ou pergunta está vazio. Por favor, forneça ambos."

        # Função que chama 'answer_question' com tradução
        def process_question():
            context = context_text.get("1.0", tk.END).strip()
            question = question_text.get("1.0", tk.END).strip()

            # Traduz a pergunta do português para o inglês
            translated_question = translate_text_model(question, "pt", "en")

            # Obtém a resposta do modelo em inglês
            english_answer = answer_question(context, translated_question)

            # Traduz a resposta do inglês para o português
            translated_answer = translate_text_model(english_answer, "en", "pt")

            # Exibe a resposta traduzida na interface
            answer_text.delete("1.0", tk.END)
            answer_text.insert(tk.END, translated_answer)

        # Botão para processar a pergunta
        answer_button = customtkinter.CTkButton(tab, text="Responder", command=process_question)
        answer_button.grid(row=3, column=1, padx=20, pady=10)

    
    def create_translation_tab(self, tab):
        """Cria os elementos para a aba de tradução."""
        tab.grid_rowconfigure((0, 1, 2, 3, 4), weight=1)
        tab.grid_columnconfigure((0, 1), weight=1)

        # Campo de entrada para o texto a ser traduzido
        input_label = customtkinter.CTkLabel(tab, text="Texto a ser traduzido:")
        input_label.grid(row=0, column=0, padx=10, pady=10, sticky="w")

        self.translation_input_text = customtkinter.CTkTextbox(tab, height=150)
        self.translation_input_text.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

        # Seleção de idioma de origem e destino
        from_label = customtkinter.CTkLabel(tab, text="Idioma de origem:")
        from_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        self.from_language_option = customtkinter.CTkOptionMenu(tab, values=["Inglês", "Português", "Espanhol"])
        self.from_language_option.set("Português")  # Idioma padrão
        self.from_language_option.grid(row=2, column=1, padx=10, pady=5, sticky="w")

        to_label = customtkinter.CTkLabel(tab, text="Idioma de destino:")
        to_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        self.to_language_option = customtkinter.CTkOptionMenu(tab, values=["Inglês", "Português", "Espanhol"])
        self.to_language_option.set("Inglês")  # Idioma padrão
        self.to_language_option.grid(row=3, column=1, padx=10, pady=5, sticky="w")

        # Botão para traduzir o texto
        translate_button = customtkinter.CTkButton(tab, text="Traduzir", command=self.translate_text)
        translate_button.grid(row=4, column=0, columnspan=2, padx=10, pady=10)

        # Campo de saída para o texto traduzido
        output_label = customtkinter.CTkLabel(tab, text="Texto traduzido:")
        output_label.grid(row=5, column=0, padx=10, pady=10, sticky="w")

        self.translation_output_text = customtkinter.CTkTextbox(tab, height=150)
        self.translation_output_text.grid(row=6, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

    def translate_text(self):
        """Traduz o texto de um idioma para outro usando Argos Translate."""
        try:
            # Mapeamento de idiomas para códigos ISO 639-1
            language_map = {
                "Inglês": "en",
                "Espanhol": "es",
                "Português": "pt"
            }

            # Obtém os textos de entrada e os códigos de idioma
            text_to_translate = self.translation_input_text.get("1.0", tk.END).strip()
            from_lang = language_map[self.from_language_option.get()]
            to_lang = language_map[self.to_language_option.get()]

            if not text_to_translate:
                messagebox.showwarning("Aviso", "Por favor, insira um texto para traduzir.")
                return

            # Atualiza os pacotes disponíveis
            argostranslate.package.update_package_index()
            available_packages = argostranslate.package.get_available_packages()

            # Tradução direta
            package_to_install = next(
                (pkg for pkg in available_packages if pkg.from_code == from_lang and pkg.to_code == to_lang),
                None
            )

            if package_to_install:
                argostranslate.package.install_from_path(package_to_install.download())
                translated_text = argostranslate.translate.translate(text_to_translate, from_lang, to_lang)
            else:
                messagebox.showwarning("Aviso", f"Nenhum pacote disponível para traduzir de {self.from_language_option.get()} para {self.to_language_option.get()}.")
                return

            # Exibe o texto traduzido
            self.translation_output_text.delete("1.0", tk.END)
            self.translation_output_text.insert(tk.END, translated_text)

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao traduzir o texto: {str(e)}")
    
    def process_abstracts(self, metadata_list):
        """Processa os abstracts dos metadados e responde as perguntas pré-definidas."""
        results = {}  # Dicionário para armazenar as respostas para cada artigo

        for i, metadata in enumerate(metadata_list):
            abstract = metadata.get('Abstract', '')
            if not abstract:
                # Se não houver abstract, criamos um dicionário com uma mensagem de erro para cada pergunta
                article_results = {
                    question: "No abstract available to answer this question."
                    for question in PREDEFINED_QUESTIONS
                }
            else:
                # Se houver abstract, respondemos as perguntas
                article_results = {}
                for question in PREDEFINED_QUESTIONS:
                    answer = question_answerer(question=question, context=abstract)  # Usa o modelo aqui
                    article_results[question] = answer

            results[f"Article {i+1}"] = article_results

        return results
        
    def answer_question(self, question, context):
        """Responde uma pergunta com base no contexto usando o modelo."""
        try:
            # Chama o modelo para responder a pergunta
            result = self.question_answerer(question=question, context=context)
            
            # Define um limite de confiança (score > 0.5)
            if result["score"] > 0.5:
                return result["answer"]
            else:
                return "Não foi possível encontrar uma resposta clara no abstract."
        except Exception as e:
            return f"Erro ao processar a pergunta: {str(e)}"

    def format_qa_results(self, results):
        """Formata as perguntas e respostas em um dicionário."""
        formatted_results = {}
        for article, answers in results.items():
            formatted_results[article] = {}
            for question, answer in answers.items():
                if isinstance(answer, dict) and "answer" in answer:
                    formatted_results[article][question] = answer["answer"]
                else:
                    formatted_results[article][question] = str(answer)
        return formatted_results
    
    
    def generate_exclusion_criteria_chart(self, doc):
        """Gera um gráfico de barras para critérios de exclusão (apenas para artigos rejeitados)."""
        # Filtra apenas os critérios de exclusão para artigos rejeitados
        exclusion_criteria = [decision['exclusion'] for decision in self.decisions_data if decision['decision'] == "Rejeitar"]
        
        if not exclusion_criteria:  # Se não houver critérios de exclusão
            doc.add_heading("Gráfico dos Critérios de Exclusão", level=2)
            doc.add_paragraph("Nenhum artigo foi rejeitado.")
            return

        # Conta a frequência dos critérios de exclusão
        criteria_counts = Counter(exclusion_criteria)

        # Cria o gráfico de barras
        plt.figure(figsize=(8, 4))
        sns.barplot(x=list(criteria_counts.keys()), y=list(criteria_counts.values()))
        plt.title("Critérios de Exclusão")
        plt.xlabel("Critério")
        plt.ylabel("Número de Estudos")
        plt.xticks(rotation=45)

        # Salva o gráfico em um buffer de memória
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)

        # Adiciona o gráfico ao documento
        doc.add_heading("Gráfico dos Critérios de Exclusão", level=2)
        doc.add_picture(buf, width=Inches(6))
        plt.close()

    def generate_exclusion_criteria_chart_2(self, output_path):
        """
        Gera um gráfico de barras para critérios de exclusão, considerando apenas artigos rejeitados,
        e salva como arquivo de imagem.
        """
        # Filtra apenas critérios de exclusão de artigos rejeitados
        exclusion_criteria = [decision['exclusion'] for decision in self.decisions_data if decision['decision'] == "Rejeitar"]

        if not exclusion_criteria:
            print("Nenhum artigo foi rejeitado. Gráfico não será gerado.")
            return

        # Conta a frequência dos critérios de exclusão
        criteria_counts = Counter(exclusion_criteria)

        # Cria o gráfico de barras
        plt.figure(figsize=(8, 4))
        sns.barplot(x=list(criteria_counts.keys()), y=list(criteria_counts.values()))
        plt.title("Critérios de Exclusão (Artigos Rejeitados)")
        plt.xlabel("Critério")
        plt.ylabel("Número de Estudos")
        plt.xticks(rotation=45)

        # Salva o gráfico em um arquivo de imagem
        plt.savefig(output_path, format='png', bbox_inches='tight')
        plt.close()

    def generate_priority_pie_chart(self, doc):
        """Gera um gráfico de pizza para prioridades de leitura (apenas para artigos aceitos)."""
        # Filtra apenas as prioridades para artigos aceitos
        priorities = [decision['priority'] for decision in self.decisions_data if decision['decision'] == "Aceitar"]
        
        if not priorities:  # Se não houver prioridades
            doc.add_heading("Prioridades de Leitura", level=2)
            doc.add_paragraph("Nenhum artigo foi aceito.")
            return

        # Conta a frequência das prioridades
        priority_counts = Counter(priorities)

        # Cria o gráfico de pizza
        plt.figure(figsize=(6, 6))
        plt.pie(priority_counts.values(), labels=priority_counts.keys(), autopct='%1.1f%%')
        plt.title("Distribuição de Prioridades de Leitura")

        # Salva o gráfico em um buffer de memória
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)

        # Adiciona o gráfico ao documento
        doc.add_heading("Prioridades de Leitura", level=2)
        doc.add_picture(buf, width=Inches(4))
        plt.close()

    def generate_priority_pie_chart_2(self, output_path):
        """
        Gera um gráfico de pizza para prioridades de leitura, considerando apenas artigos aceitos,
        e salva como arquivo de imagem.
        """
        # Filtra apenas prioridades de artigos aceitos
        priorities = [decision['priority'] for decision in self.decisions_data if decision['decision'] == "Aceitar"]

        if not priorities:
            print("Nenhum artigo foi aceito. Gráfico não será gerado.")
            return

        # Conta a frequência das prioridades
        priority_counts = Counter(priorities)

        # Cria o gráfico de pizza
        plt.figure(figsize=(6, 6))
        plt.pie(priority_counts.values(), labels=priority_counts.keys(), autopct='%1.1f%%')
        plt.title("Distribuição de Prioridades de Leitura (Artigos Aceitos)")

        # Salva o gráfico em um arquivo de imagem
        plt.savefig(output_path, format='png', bbox_inches='tight')
        plt.close()
    
        

    def save_to_docx(self):
        """Salva os dados da revisão sistemática em um arquivo .docx"""
        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                filetypes=[("Documentos Word", "*.docx")])
        if not file_path:
            return
        
        doc = Document()
        doc.add_heading("Revisão Sistemática", level=1)
        
        # Adiciona título da revisão
        doc.add_heading("Título da Revisão", level=2)
        title_paragraph = doc.add_paragraph(self.title_entry.get())
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica o parágrafo

        # Adiciona pesquisadores
        doc.add_heading("Pesquisadores", level=2)
        researchers_paragraph = doc.add_paragraph(self.researchers_entry.get())
        researchers_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica o parágrafo

        # Adiciona descrição
        doc.add_heading("Descrição", level=2)
        description_paragraph = doc.add_paragraph(self.description_textbox.get("1.0", "end").strip())
        description_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica o parágrafo

        # Adiciona objetivo
        doc.add_heading("Objetivo", level=2)
        objective_paragraph = doc.add_paragraph(self.objective_entry.get())
        objective_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Adiciona linguagem dos artigos
        doc.add_heading("Linguagem dos Artigos", level=2)
        selected_languages = [var.get() for _, var in self.language_checkboxes if var.get()]
        doc.add_paragraph(", ".join(selected_languages))

        # Adiciona editoras
        doc.add_heading("Editoras", level=2)
        selected_publishers = [var.get() for _, var in self.publisher_checkboxes if var.get()]
        doc.add_paragraph(", ".join(selected_publishers))

        # Adiciona palavras-chave
        doc.add_heading("Palavras-chave", level=2)
        key_paragraph = doc.add_paragraph(self.keywords_textbox.get())
        key_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Adiciona as strings de pesquisa
        doc.add_heading("Strings de pesquisa", level=2)
        string_paragraph = doc.add_paragraph(self.strings_textbox.get())
        string_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Adiciona critérios de seleção e exclusão
        doc.add_heading("Critérios de Seleção", level=2)
        criteria_selection_paragraph = doc.add_paragraph(self.criteria_textbox_s.get())
        criteria_selection_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   # Justifica o parágrafo

        doc.add_heading("Critérios de Exclusão", level=2)
        criteria_exclusion_paragraph = doc.add_paragraph(self.criteria_textbox_e.get())
        criteria_exclusion_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica o parágrafo

        # Adiciona os metadados dos artigos analisados
        for metadata in self.metadata_dados:
            # Adiciona o título
            title = metadata.get('Title', 'N/A')
            p = doc.add_paragraph()
            p.add_run('Título: ').bold = True
            p.add_run(title)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona os autores
            authors = metadata.get('Authors', [])
            p = doc.add_paragraph()
            p.add_run('Autores: ').bold = True
            p.add_run(', '.join(authors))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona a revista
            journal = metadata.get('Journal', 'N/A')
            p = doc.add_paragraph()
            p.add_run('Revista: ').bold = True
            p.add_run(journal)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona o DOI
            doi = metadata.get('DOI', 'N/A')
            p = doc.add_paragraph()
            p.add_run('DOI: ').bold = True
            p.add_run(doi)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona o resumo
            abstract = metadata.get('Abstract', 'N/A')
            p = doc.add_paragraph()
            p.add_run('Resumo: ').bold = True
            p.add_run(abstract)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona o ano
            year = metadata.get('Year', 'N/A')
            p = doc.add_paragraph()
            p.add_run('Ano: ').bold = True
            p.add_run(year)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Adiciona uma linha em branco entre os artigos
            doc.add_paragraph()

        
        # Adiciona as perguntas e respostas de cada artigo
        doc.add_heading("Perguntas e respostas dos artigos", level=2)
        doc.add_paragraph()  # Adiciona um espaço em branco

        if hasattr(self, 'formatted_qa') and isinstance(self.formatted_qa, dict):
            for article, qa_pairs in self.formatted_qa.items():
                    # Adiciona o título do artigo
                doc.add_heading(article, level=3)

                # Adiciona cada pergunta e resposta
                for question, answer in qa_pairs.items():
                    # Adiciona a pergunta
                    question_paragraph = doc.add_paragraph()
                    question_paragraph.add_run("Pergunta: ").bold = True
                    question_paragraph.add_run(question)
                    question_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Adiciona a resposta
                    answer_paragraph = doc.add_paragraph()
                    answer_paragraph.add_run("Resposta: ").bold = True
                    answer_paragraph.add_run(str(answer))  # Converte a resposta para string
                    answer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Adiciona um espaço entre as perguntas e respostas
                    doc.add_paragraph()
        else:
            # Se não houver perguntas e respostas, adiciona uma mensagem
            doc.add_paragraph("Nenhuma pergunta e resposta disponível.")
        
        
        """Adiciona uma tabela ao documento com as decisões dos artigos."""
        doc.add_heading("Decisão de seleção dos artigos", level=2)
        table = doc.add_table(rows=1, cols=5)  # Adiciona uma coluna extra para prioridade
        table.style = 'Table Grid'

        # Cabeçalho da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Artigo'
        hdr_cells[1].text = 'Critério de seleção'
        hdr_cells[2].text = 'Critério de exclusão'
        hdr_cells[3].text = 'Decisão'
        hdr_cells[4].text = 'Prioridade de leitura'

        # Adiciona os dados dos artigos à tabela
        if hasattr(self, 'decisions_data') and self.decisions_data:
            for i, decision in enumerate(self.decisions_data):
                row_cells = table.add_row().cells
                row_cells[0].text = f"Artigo {i+1}"
                row_cells[1].text = decision['selection'] if decision['decision'] == "Aceitar" else "N/A"
                row_cells[2].text = decision['exclusion'] if decision['decision'] == "Rejeitar" else "N/A"
                row_cells[3].text = decision['decision']
                row_cells[4].text = decision['priority'] if decision['decision'] == "Aceitar" else "N/A"


        self.generate_exclusion_criteria_chart(doc)
        self.generate_priority_pie_chart(doc)

        # Adiciona as referências ao DOCX
        doc.add_heading("Referências", level=2)
        doc.add_heading("", level=2)

        if self.metadata_ref:  # Verifica se há referências formatadas
            for reference in self.metadata_ref:
                p = doc.add_paragraph(reference)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justifica o parágrafo
        else:
            doc.add_paragraph("Nenhum artigo analisado.")


        # Salva o arquivo
        doc.save(file_path)

    def load_from_docx(self):
        """Carrega os dados de um arquivo .docx para os campos da interface"""
        file_path = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if not file_path:
            return

        doc = Document(file_path)
        sections = {}

        current_section = None
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                current_section = para.text.strip()
                sections[current_section] = []
            elif current_section:
                sections[current_section].append(para.text.strip())

        # Preenchendo os campos
        self.title_entry.delete(0, "end")
        self.title_entry.insert(0, sections.get("Título da Revisão", [""])[0])

        self.researchers_entry.delete(0, "end")
        self.researchers_entry.insert(0, sections.get("Pesquisadores", [""])[0])

        self.description_textbox.delete("1.0", "end")
        self.description_textbox.insert("1.0", "\n".join(sections.get("Descrição", [""])) or "")

        self.objective_entry.delete(0, "end")
        self.objective_entry.insert(0, sections.get("Objetivo", [""])[0])

        # Corrigindo a leitura de Linguagens
        languages_str = sections.get("Linguagem dos Artigos", [""])[0]
        selected_languages = [lang.strip() for lang in languages_str.split(",") if lang.strip()]
        for checkbox, var in self.language_checkboxes:
            if checkbox.cget("text") in selected_languages:
                checkbox.select()
            else:
                checkbox.deselect()

        # Corrigindo a leitura de Editoras
        publishers_str = sections.get("Editoras", [""])[0]
        selected_publishers = [pub.strip() for pub in publishers_str.split(",") if pub.strip()]
        for checkbox, var in self.publisher_checkboxes:
            if checkbox.cget("text") in selected_publishers:
                checkbox.select()
            else:
                checkbox.deselect()

        self.keywords_textbox.delete(0, "end")
        self.keywords_textbox.insert(0, sections.get("Palavras-chave", [""])[0])

        self.strings_textbox.delete(0, "end")
        self.strings_textbox.insert(0, sections.get("Strings de pesquisa", [""])[0])

        self.criteria_textbox_s.delete(0, "end")
        self.criteria_textbox_s.insert(0, sections.get("Critérios de Seleção", [""])[0])

        self.criteria_textbox_e.delete(0, "end")
        self.criteria_textbox_e.insert(0, sections.get("Critérios de Exclusão", [""])[0])

    def format_metadata_for_latex(self, metadata):
        """Formata os metadados para o LaTeX, evitando indentação automática."""
        formatted_text = ""
        for i, article in enumerate(metadata):
            formatted_text += f"\\subsection*{{Artigo {i + 1}}}\n"
            
            # Título
            formatted_text += "\\noindent\\textbf{Título}: " + self.escape_latex_special_chars(self._to_string(article.get('Title', 'N/A'))) + "\n\n"
            
            # Autores (pode ser uma lista)
            authors = article.get('Authors', 'N/A')
            if isinstance(authors, list):
                authors = ", ".join(authors)  # Converte a lista em uma string separada por vírgulas
            formatted_text += "\\noindent\\textbf{Autores}: " + self.escape_latex_special_chars(authors) + "\n\n"
            
            # Revista
            formatted_text += "\\noindent\\textbf{Revista}: " + self.escape_latex_special_chars(self._to_string(article.get('Journal', 'N/A'))) + "\n\n"
            
            # DOI
            formatted_text += "\\noindent\\textbf{DOI}: " + self.escape_latex_special_chars(self._to_string(article.get('DOI', 'N/A'))) + "\n\n"
            
            # Ano
            formatted_text += "\\noindent\\textbf{Ano}: " + self.escape_latex_special_chars(self._to_string(article.get('Year', 'N/A'))) + "\n\n"
            
            # Resumo
            formatted_text += "\\noindent\\textbf{Resumo}: " + self.escape_latex_special_chars(self._to_string(article.get('Abstract', 'N/A'))) + "\n\n"
            
            # Linha horizontal para separar artigos
            #formatted_text += "\\hrule\n\n"
        
        return formatted_text

    def _to_string(self, value):
        """Converte um valor para string, lidando com listas."""
        if isinstance(value, list):
            return ", ".join(value)  # Converte a lista em uma string separada por vírgulas
        return str(value)  # Converte outros tipos para string
        
    def escape_latex_special_chars(self, text):
        """Escapa caracteres especiais para o LaTeX."""
        special_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '^': '\\textasciicircum{}',
            '\\': '\\textbackslash{}',
        }
        for char, escaped_char in special_chars.items():
            text = text.replace(char, escaped_char)
        return text
    
    def format_qa_results_for_latex(self, results):
        """Formata as perguntas e respostas para o LaTeX."""
        if not isinstance(results, dict):
            return "\\section{Perguntas e Respostas dos Artigos}\nNenhuma pergunta e resposta disponível.\n\n"

        formatted_text = "\\section{Perguntas e Respostas dos Artigos}\n"
        for article, answers in results.items():
            formatted_text += f"\\subsection*{{{article}}}\n"
            for question, answer in answers.items():
                formatted_text += f"\\noindent\\textbf{{Pergunta}}: {self.escape_latex_special_chars(question)}\n\n"
                formatted_text += f"\\noindent\\textbf{{Resposta}}: {self.escape_latex_special_chars(answer)}\n\n"
            formatted_text += "\n\n"  # Linha em branco para separar artigos
        return formatted_text
    
    def format_references_for_latex(self, references):
        formatted_text = "\\section{Referências}\n"
        formatted_text += "\\begin{enumerate}\n"
        for ref in references:
            formatted_text += f"    \\item {ref}\n"
        formatted_text += "\\end{enumerate}\n"
        return formatted_text

    
    def save_to_latex(self):
        # Solicita ao usuário o local para salvar o arquivo .tex
        file_path = filedialog.asksaveasfilename(defaultextension=".tex",
                                                filetypes=[("Arquivo LaTeX", "*.tex")])
        if not file_path:
            return

        # Cria uma pasta com o mesmo nome do arquivo .tex (sem a extensão)
        folder_path = os.path.splitext(file_path)[0]  # Remove a extensão .tex
        os.makedirs(folder_path, exist_ok=True)  # Cria a pasta se não existir

        # Caminhos para salvar os gráficos dentro da pasta
        exclusion_chart_path = os.path.join(folder_path, "exclusion_criteria.png")
        priority_chart_path = os.path.join(folder_path, "priority_pie_chart.png")

        # Gera os gráficos e salva como imagens na pasta
        self.generate_exclusion_criteria_chart_2(exclusion_chart_path)
        self.generate_priority_pie_chart_2(priority_chart_path)

        # Início do documento LaTeX com pacotes e configurações
        latex_content = "\\documentclass{article}\n"
        latex_content += "\\usepackage[utf8]{inputenc}\n"
        latex_content += "\\usepackage{amsmath}\n"
        latex_content += "\\usepackage{hyperref}\n"  # Para links clicáveis
        latex_content += "\\usepackage{booktabs}\n"  # Para tabelas profissionais
        latex_content += "\\usepackage{geometry}\n"  # Para ajustar margens
        latex_content += "\\usepackage{enumitem}\n"  # Para listas personalizadas
        latex_content += "\\usepackage{graphicx}\n"  # Para incluir imagens
        latex_content += "\\geometry{a4paper, margin=1in}\n"  # Configuração de margens
        latex_content += "\\begin{document}\n"

        # Título e autores
        latex_content += "\\title{" + self.title_entry.get() + "}\n"
        latex_content += "\\author{" + self.researchers_entry.get() + "}\n"
        latex_content += "\\maketitle\n"

        # Função auxiliar para adicionar seções
        def add_section(title, content):
            if content.strip():  # Verifica se o conteúdo não está vazio
                return f"\\section{{{title}}}\n{content}\n\n"
            return f"\\section{{{title}}}\nNenhum conteúdo disponível.\n\n"

        # Adicionando as seções do LaTeX
        latex_content += add_section("Descrição", self.description_textbox.get("1.0", "end").strip())
        latex_content += add_section("Objetivo", self.objective_entry.get())
        latex_content += add_section("Linguagem dos Artigos", ", ".join([var.get() for _, var in self.language_checkboxes if var.get()]))
        latex_content += add_section("Editoras", ", ".join([var.get() for _, var in self.publisher_checkboxes if var.get()]))
        latex_content += add_section("Palavras-chave", self.keywords_textbox.get())
        latex_content += add_section("Strings de Pesquisa", self.strings_textbox.get())
        latex_content += add_section("Critérios de Seleção", self.criteria_textbox_s.get())
        latex_content += add_section("Critérios de Exclusão", self.criteria_textbox_e.get())

        
        # Adicionando os metadados dos artigos analisados
        latex_content += "\\section{Artigos Analisados}\n"
        if hasattr(self, 'clean_metadata') and self.clean_metadata:
            latex_content += self.format_metadata_for_latex(self.clean_metadata) + "\n\n"
        else:
            latex_content += "Nenhum artigo analisado.\n\n"

        # Adicionando as perguntas e respostas dos artigos
        #latex_content += "\\section{Perguntas e Respostas dos Artigos}\n"
        if hasattr(self, 'formatted_qa') and isinstance(self.formatted_qa, dict):
            latex_content += self.format_qa_results_for_latex(self.formatted_qa) + "\n\n"
        else:
            latex_content += "Nenhuma pergunta e resposta disponível.\n\n"

        # Adicionando tabelas
        latex_content += "\\section{Decisão de Seleção dos Artigos}\n"

        if hasattr(self, 'decisions_data') and self.decisions_data:
            latex_content += "\\begin{tabular}{|p{1.0cm}|p{4.8cm}|p{4.8cm}|p{1.2cm}|p{1.8cm}|}\n"
            latex_content += "\\hline\n"
            latex_content += "Artigo & Critério de Seleção & Critério de Exclusão & Decisão & Prioridade de Leitura \\\\\n"
            latex_content += "\\hline\n"
            
            for i, decision in enumerate(self.decisions_data):
                # Trata os valores com base na decisão
                selection = decision['selection'] if decision['decision'] == "Aceitar" else "N/A"
                exclusion = decision['exclusion'] if decision['decision'] == "Rejeitar" else "N/A"
                priority = decision['priority'] if decision['decision'] == "Aceitar" else "N/A"
                
                # Adiciona a linha da tabela
                latex_content += f"{i+1} & {selection} & {exclusion} & {decision['decision']} & {priority} \\\\\n"
                latex_content += "\\hline\n"
            
            latex_content += "\\end{tabular}\n\n"
        else:
            latex_content += "Nenhuma decisão de seleção disponível.\n\n"

        # Adicionando os gráficos
        latex_content += "\\section{Critérios de Exclusão}\n"
        latex_content += f"\\includegraphics[width=\\textwidth]{{{os.path.basename(exclusion_chart_path)}}}\n\n"

        latex_content += "\\section{Prioridades de Leitura}\n"
        latex_content += f"\\includegraphics[width=0.6\\textwidth]{{{os.path.basename(priority_chart_path)}}}\n\n"

        # Adicionando as referências
        if hasattr(self, 'metadata_ref') and self.metadata_ref:
            latex_content += self.format_references_for_latex(self.metadata_ref) + "\n\n"
        else:
            latex_content += "Nenhuma referência encontrada.\n\n"

        # Fim do documento LaTeX
        latex_content += "\\end{document}\n"

        # Salva o arquivo LaTeX na pasta
        tex_file_path = os.path.join(folder_path, os.path.basename(file_path))
        with open(tex_file_path, 'w', encoding='utf-8') as f:
            f.write(latex_content)

        # Cria um arquivo ZIP contendo a pasta
        zip_path = folder_path + ".zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, start=folder_path)
                    zipf.write(file_path, arcname)

        print(f"Arquivo LaTeX e imagens salvos em: {folder_path}")
        print(f"Pasta zipada criada em: {zip_path}")
    
    def load_from_latex(self):
        """Carrega os dados de um arquivo .tex para os campos da interface"""
        file_path = filedialog.askopenfilename(filetypes=[("Arquivo LaTeX", "*.tex")])
        if not file_path:
            return

        with open(file_path, 'r', encoding='utf-8') as f:
            latex_content = f.read()

        # Extrai as seções do arquivo LaTeX
        sections = {}
        section_titles = ["Descrição", "Objetivo",
                        "Linguagem dos Artigos", "Editoras", "Palavras-chave", 
                        "Strings de Pesquisa", "Critérios de Seleção",
                        "Critérios de Exclusão"]

        # Extrai o título do comando \title{}
        title_start = latex_content.find("\\title{") + len("\\title{")
        if title_start >= len("\\title{"):
            title_end = latex_content.find("}", title_start)
            sections["Título da Revisão"] = latex_content[title_start:title_end].strip()

        for title in section_titles:
            # Tenta encontrar a seção com e sem asterisco
            for marker in [f"\\section{{{title}}}", f"\\section*{{{title}}}"]:
                if marker in latex_content:
                    start_idx = latex_content.find(marker) + len(marker)
                    end_idx = latex_content.find("\\section", start_idx)
                    if end_idx == -1:
                        end_idx = len(latex_content)
                    section_content = latex_content[start_idx:end_idx].strip()
                    # Remove comandos LaTeX comuns preservando o conteúdo
                    section_content = section_content.replace("\\textbf{", "").replace("}", "")
                    section_content = section_content.replace("\\textit{", "").replace("\\emph{", "")
                    sections[title] = section_content
                    break

        # Preenchendo os campos com as seções extraídas
        self.title_entry.delete(0, "end")
        self.title_entry.insert(0, sections.get("Título da Revisão", ""))

        self.researchers_entry.delete(0, "end")
        # Pesquisadores podem ser extraídos do comando \author{}
        author_start = latex_content.find("\\author{") + len("\\author{")
        if author_start >= len("\\author{"):
            author_end = latex_content.find("}", author_start)
            self.researchers_entry.insert(0, latex_content[author_start:author_end].strip())

        self.description_textbox.delete("1.0", "end")
        self.description_textbox.insert("1.0", sections.get("Descrição", ""))

        self.objective_entry.delete(0, "end")
        self.objective_entry.insert(0, sections.get("Objetivo", ""))

        # Linguagens dos artigos
        languages_str = sections.get("Linguagem dos Artigos", "")
        selected_languages = [lang.strip() for lang in languages_str.split(",") if lang.strip()]
        for checkbox, var in self.language_checkboxes:
            if checkbox.cget("text") in selected_languages:
                checkbox.select()
            else:
                checkbox.deselect()

        # Editoras 
        publishers_str = sections.get("Editoras", "")
        selected_publishers = [pub.strip() for pub in publishers_str.split(",") if pub.strip()]
        for checkbox, var in self.publisher_checkboxes:
            if checkbox.cget("text") in selected_publishers:
                checkbox.select()
            else:
                checkbox.deselect()

        self.keywords_textbox.delete(0, "end")
        self.keywords_textbox.insert(0, sections.get("Palavras-chave", ""))

        self.strings_textbox.delete(0, "end")
        self.strings_textbox.insert(0, sections.get("Strings de Pesquisa", ""))

        self.criteria_textbox_s.delete(0, "end")
        self.criteria_textbox_s.insert(0, sections.get("Critérios de Seleção", ""))

        self.criteria_textbox_e.delete(0, "end")
        self.criteria_textbox_e.insert(0, sections.get("Critérios de Exclusão", ""))
    


customtkinter.set_default_color_theme("green")  # Themes: "blue" (standard), "green", "dark-blue"
if __name__ == "__main__":
    app = App()
    app.collect_decisions()
    app.mainloop()

